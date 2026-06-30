from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from src.models.profile import Profile
from src.styles.theme import Theme


class HeaderComponent:

    def __init__(self, document, profile: Profile):
        self.document = document
        self.profile = profile

    def build(self):

        # Name
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run(self.profile.full_name.upper())
        run.bold = True
        run.font.name = Theme.FONT
        run.font.size = Theme.NAME_SIZE
        run.font.color.rgb = Theme.PRIMARY

        # Title
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run(self.profile.professional_title)
        run.font.name = Theme.FONT
        run.font.size = Theme.TITLE_SIZE
        run.font.color.rgb = Theme.SECONDARY

        # Contact Details
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run(
            f"{self.profile.email} | "
            f"{self.profile.phone} | "
            f"{self.profile.location}"
        )

        run.font.name = Theme.FONT
        run.font.size = Theme.BODY_SIZE